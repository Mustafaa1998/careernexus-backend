from __future__ import annotations
import os
import logging
from typing import Optional, List
import re

import pdfplumber
from docx import Document

from .utils import get_extension, clean_whitespace, safe_slice_lines
from .preprocessing import normalize_text, split_sections
from .ocr import pdf_to_text_via_ocr
from .extractors import (
    extract_email, extract_phone, extract_name, extract_years,
    parse_education, parse_experience, extract_languages, extract_summary,
    extract_certifications, load_skills_dictionary, extract_skills,
    load_gazetteer, extract_location,
)
from .schema import ParseResponse, EducationItem, ExperienceItem
from .config import MIN_SELECTABLE_TEXT_CHARS, OCR_ENABLED

logger = logging.getLogger("resume_parser.parser")

RESOURCES_DIR = os.path.join(os.path.dirname(__file__), "resources")
SKILLS_PATH = os.path.join(RESOURCES_DIR, "skills.json")
GAZETTEER_PATH = os.path.join(RESOURCES_DIR, "location_gazetteer.txt")

# ========== Column Detection ==========
def detect_columns(pdf_path: str) -> int:
    """
    Detects the number of columns in the PDF file by analyzing the first page.
    Returns:
        1 if single-column layout
        2 if two-column layout
    """
    with pdfplumber.open(pdf_path) as pdf:
        first_page = pdf.pages[0]
        # Extracting character positions (x0, x1 values indicate column boundaries)
        widths = [char['doctop'] for char in first_page.chars]
        min_width = min(widths)
        max_width = max(widths)
        average_width = (max_width - min_width) / len(set(widths))  # Calculate average spacing

        # Heuristic: If columns are wide apart, it's a two-column layout
        if average_width > 100:  # Threshold for column separation
            return 2
        return 1

# ========== Read PDF (Selectable Text) ==========
def read_pdf_text(path: str) -> str:
    """Extract selectable text from a PDF via pdfplumber (no OCR)."""
    parts: List[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    parts.append(t)
    except Exception as e:
        logger.warning("pdfplumber failed on %s: %s", path, e)
    return "\n\n".join(parts)

# ========== Read DOCX ==========
def read_docx_text(path: str) -> str:
    """Extract text (and tables) from a DOCX."""
    doc = Document(path)
    parts: List[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # include tables (common in resumes)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [cell.text for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)

# ========== Crop Columns Based on Layout ==========
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

def crop_columns(pdf_path: str, column_count: int = 1) -> str:
    """Crop columns if multiple columns are detected in the PDF, else use the full page."""
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[0]
        if column_count == 1:
            return page.extract_text()
        
        # If 2 columns, split the page into left and right
        width = page.width
        left_bbox = (0, 0, width / 2, page.height)
        right_bbox = (width / 2, 0, width, page.height)

        left_image = page.to_image().crop(left_bbox)
        right_image = page.to_image().crop(right_bbox)

        # OCR for left and right columns separately if text is not selectable
        left_text = pytesseract.image_to_string(left_image)
        right_text = pytesseract.image_to_string(right_image)
        
        return left_text + "\n" + right_text

# ========== Section Extraction (Regex) ==========
def extract_section(text: str, section_name: str) -> str:
    """Extract a specific section based on keywords like 'Summary', 'Experience'."""
    section_pattern = r"(?i)(?<=\b" + section_name + r"\b)(.*?)(?=\n\S)"
    match = re.search(section_pattern, text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return ""

def extract_all_sections(text: str):
    """Extract key sections (summary, experience, education, etc.) from resume text."""
    summary = extract_section(text, "summary")
    experience = extract_section(text, "experience")
    education = extract_section(text, "education")
    projects = extract_section(text, "projects")
    skills = extract_section(text, "skills")
    certifications = extract_section(text, "certifications")
    return {
        "summary": summary,
        "experience": experience,
        "education": education,
        "projects": projects,
        "skills": skills,
        "certifications": certifications,
    }

# ========== Convert Text to JSON ==========
def output_json(parsed_data: dict, output_file: str):
    """Save the extracted data as JSON to file."""
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(parsed_data, f, ensure_ascii=False, indent=4)


import json

def output_json(parsed_data: dict, output_file: str):
    """Save the extracted data as JSON to file."""
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(parsed_data, f, ensure_ascii=False, indent=4)
        logger.info(f"Successfully saved parsed data to {output_file}")
    except Exception as e:
        logger.error(f"Error while saving JSON to {output_file}: {e}")

# ========== Main Parsing Function ==========
def parse_text_to_json(text: str, original_filename: Optional[str] = None) -> ParseResponse:
    text_clean = normalize_text(clean_whitespace(text or ""))
    logger.info("Parsing text preview: %s", safe_slice_lines(text_clean, 500))

    # Field extractors
    name = extract_name(text_clean)
    email = extract_email(text_clean)
    phone = extract_phone(text_clean)
    years = extract_years(text_clean)
    summary = extract_summary(text_clean)
    languages = extract_languages(text_clean)
    certifications = extract_certifications(text_clean)

    # Sections
    sections = extract_all_sections(text_clean)
    edu = parse_education(sections.get("education", ""))
    exp = parse_experience(sections.get("experience", sections.get("work experience", "")))

    # Skills via dictionary + fuzzy match
    skills_dict = load_skills_dictionary(SKILLS_PATH)
    skills = extract_skills(text_clean, skills_dict)

    # Location via gazetteer (looks near header/contact section); tolerate missing file
    try:
        gazetteer = load_gazetteer(GAZETTEER_PATH)
    except Exception:
        gazetteer = []
    location = extract_location(text_clean, gazetteer) if gazetteer else None

    # Pack dataclasses
    edu_items = [EducationItem(degree=e["degree"], institution=e["institution"], year=e["year"]) for e in edu]
    exp_items = [ExperienceItem(**e) for e in exp]

    return ParseResponse(
        filename=original_filename,
        name=name,
        email=email,
        phone=phone,
        location=location,
        summary=summary,
        education=edu_items,
        experience=exp_items,
        skills=skills,
        certifications=certifications,
        languages=languages,
        years=years,
        raw_text=text_clean[:100000],  # cap to protect API payloads
    )


def process_resume(pdf_path: str, output_file: str):
    """Process the resume: detect columns, extract text, apply OCR, and save output as JSON."""
    try:
        logger.info(f"Processing resume: {pdf_path}")

        # Step 1: Detect columns in the PDF
        column_count = detect_columns(pdf_path)
        logger.info(f"Detected column count: {column_count} columns")

        # Step 2: Extract text (with OCR if necessary)
        text = crop_columns(pdf_path, column_count)

        # Step 3: Extract sections (summary, experience, education, etc.)
        sections = extract_all_sections(text)

        # Step 4: Extract skills and other data
        parsed_data = parse_text_to_json(text, original_filename=pdf_path)

        # Step 5: Output the parsed data as JSON
        output_json(parsed_data, output_file)
        
    except Exception as e:
        logger.error(f"Error while processing the resume: {pdf_path}. Error: {e}")


def parse_file(path: str, original_filename: Optional[str] = None) -> ParseResponse:
    """Main entry: read file (PDF/DOCX), optionally OCR, then parse to structured JSON."""
    ext = get_extension(path)
    display_name = original_filename or os.path.basename(path)
    logger.info("Parsing file %s (%s)", display_name, ext)

    # Detect column count (1 or 2)
    column_count = detect_columns(path)
    logger.info("Detected column count: %s", column_count)

    if ext == ".pdf":
        # 1) Try selectable text
        text = read_pdf_text(path)
        text_len = len(text or "")
        logger.info("pdfplumber extracted %s chars", text_len)

        # 2) OCR fallback if enabled and text looks too short
        if OCR_ENABLED and (not text or text_len < MIN_SELECTABLE_TEXT_CHARS):
            logger.info(
                "Low selectable text (%s chars < %s). Using OCR fallback.",
                text_len, MIN_SELECTABLE_TEXT_CHARS
            )
            try:
                ocr_text = pdf_to_text_via_ocr(path) or ""
                if len(ocr_text) > text_len:
                    text = ocr_text
                    logger.info("OCR produced %s chars (accepted).", len(ocr_text))
                else:
                    logger.info("OCR produced %s chars (kept selectable text).", len(ocr_text))
            except Exception as e:
                logger.warning("OCR fallback failed on %s: %s", display_name, e)

        text = text or ""

    elif ext == ".docx":
        text = read_docx_text(path) or ""
    else:
        raise ValueError(f"Unsupported extension: {ext}")

    # Crop columns based on layout detection
    text = crop_columns(path, column_count)

    # Return parsed response
    return parse_text_to_json(text, original_filename=display_name)
